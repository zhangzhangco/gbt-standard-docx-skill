#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE = "标准征求意见稿修改意见表"


def load_review(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def add_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def issue_to_row(issue: dict, index: int) -> tuple[str, str, str]:
    clause_id = issue.get("clause_id") or issue.get("location_label") or issue.get("location", "")
    original = issue.get("excerpt", "")
    suggested = issue.get("suggested_action", "")
    reason_parts = [issue.get("message", "")]
    opinion = f"原文：{original}\n建议：{suggested}\n理由：{'；'.join(part for part in reason_parts if part)}"
    return str(index), clause_id, opinion


def build_report(review: dict, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    headers = ["序号", "章条\n编号", "意见"]
    widths = [1.2, 2.4, 12.0]
    header_cells = table.rows[0].cells
    for cell, text, width in zip(header_cells, headers, widths):
        add_cell_text(cell, text, bold=True, center=True)
        set_cell_width(cell, width)

    issues = review.get("issues", [])
    for idx, issue in enumerate(issues, start=1):
        row_cells = table.add_row().cells
        values = issue_to_row(issue, idx)
        add_cell_text(row_cells[0], values[0], center=True)
        add_cell_text(row_cells[1], values[1], center=True)
        add_cell_text(row_cells[2], values[2], center=False)
        for cell, width in zip(row_cells, widths):
            set_cell_width(cell, width)

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    review = load_review(Path(args.input))
    build_report(review, Path(args.output))
    print(f"已生成审查报告：{args.output}")


if __name__ == "__main__":
    main()
